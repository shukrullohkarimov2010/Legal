from flask import Flask, request, Response, jsonify, stream_with_context
from flask_cors import CORS
import requests
import json
import os
import hashlib
import time
import threading
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from pypdf import PdfReader
from zipfile import ZipFile
import xml.etree.ElementTree as ET

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# ================= НАСТРОЙКИ =================
OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "llama3.1:8b"
CHUNK_SIZE = 1500
CHUNK_OVERLAP = 200
MAX_WORKERS = 2
CACHE_FILE = ".ollama_cache.json"
TEMP_DIR = "tmp_uploads"

os.makedirs(TEMP_DIR, exist_ok=True)

# Глобальное хранилище активных сессий
active_sessions = {}
sessions_lock = threading.Lock()

# ================= КЭШИРОВАНИЕ =================
def load_cache():
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠️ Ошибка загрузки кэша: {e}")
    return {}

def save_cache(cache):
    try:
        with open(CACHE_FILE, "w", encoding="utf-8") as f:
            json.dump(cache, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"⚠️ Ошибка сохранения кэша: {e}")

def get_cache_key(text):
    return hashlib.md5((text + MODEL_NAME).encode("utf-8")).hexdigest()

# ================= ИЗВЛЕЧЕНИЕ ТЕКСТА =================
def extract_text_from_pdf(path):
    """Извлечение текста из PDF с улучшенной обработкой"""
    try:
        print(f"📄 Чтение PDF: {os.path.basename(path)}")
        reader = PdfReader(path)
        texts = []

        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text and text.strip():
                    clean_text = text.strip()
                    texts.append(clean_text)
                    print(f"  Страница {i+1}: {len(clean_text)} символов")
            except Exception as page_err:
                print(f"  ⚠️ Страница {i+1}: {page_err}")
                continue

        if not texts:
            raise Exception("PDF не содержит извлекаемого текста (возможно, сканы или изображения)")

        full_text = "\n\n".join(texts)
        print(f"  ✅ PDF: {len(full_text)} символов, {len(texts)} страниц")
        return full_text

    except Exception as e:
        raise Exception(f"Ошибка чтения PDF: {str(e)}")

def extract_text_from_docx(path):
    """Улучшенное извлечение текста из DOCX с поддержкой таблиц и резервным XML-парсингом"""
    try:
        print(f"📄 Чтение DOCX: {os.path.basename(path)}")
        doc = Document(path)
        all_text = []

        # 1. Извлекаем текст из параграфов
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                all_text.append(text)

        # 2. Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        all_text.append(text)

        # 3. Если пусто — пробуем сырой XML (резервный метод)
        if not all_text:
            print("  ⚠️ Стандартное извлечение не дало результата, пробуем XML...")
            try:
                with ZipFile(path) as z:
                    with z.open('word/document.xml') as f:
                        tree = ET.parse(f)
                        root = tree.getroot()
                        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

                        for elem in root.findall('.//w:t', ns):
                            if elem.text and elem.text.strip():
                                all_text.append(elem.text.strip())
            except Exception as xml_err:
                print(f"  ❌ XML-парсинг не удался: {xml_err}")

        if not all_text:
            raise Exception("Документ не содержит текста или повреждён")

        full_text = "\n".join(all_text)
        print(f"  ✅ DOCX: {len(full_text)} символов, {len(all_text)} блоков")
        return full_text

    except Exception as e:
        raise Exception(f"Ошибка чтения DOCX: {str(e)}")

def extract_text_from_txt(path):
    """Простое чтение TXT файлов"""
    try:
        print(f"📄 Чтение TXT: {os.path.basename(path)}")
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        if not text.strip():
            raise Exception("Файл пустой")
        print(f"  ✅ TXT: {len(text)} символов")
        return text
    except Exception as e:
        raise Exception(f"Ошибка чтения TXT: {str(e)}")

# ================= УМНАЯ РАЗБИВКА =================
def split_text_smart(text, size, overlap):
    """Разбивка текста на чанки с учётом естественных границ"""
    chunks = []
    start = 0

    while start < len(text):
        end = min(start + size, len(text))

        # Пытаемся разбить по естественным границам
        if end < len(text):
            for sep in ["\n\n", ". ", "! ", "? ", "; ", ", "]:
                pos = text.rfind(sep, start, end + overlap)
                if pos > start + size // 2:
                    end = pos + len(sep)
                    break

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # Двигаем указатель
        if end < len(text):
            start = max(start + 1, end - overlap)
        else:
            break

    return chunks if chunks else [text]  # fallback: весь текст как один чанк

# ================= ЗАПРОС К OLLAMA =================
def send_to_ollama(prompt, stream_callback=None):
    """Отправка запроса к Ollama с кэшированием и стримингом"""
    cache = load_cache()
    key = get_cache_key(prompt)

    # Проверка кэша
    if key in cache:
        cached = cache[key]
        if stream_callback:
            for char in cached:
                stream_callback(char)
                time.sleep(0.001)
        return cached

    response_text = ""
    data = {
        "model": MODEL_NAME,
        "prompt": prompt,
        "stream": True,
        "options": {
            "temperature": 0.3,
            "num_predict": 2048,
            "top_p": 0.9
        }
    }

    try:
        with requests.post(OLLAMA_URL, json=data, stream=True, timeout=600) as r:
            r.raise_for_status()
            for line in r.iter_lines():
                if line:
                    try:
                        obj = json.loads(line.decode("utf-8"))
                        chunk = obj.get("response", "")
                        if chunk:
                            response_text += chunk
                            if stream_callback:
                                stream_callback(chunk)
                    except json.JSONDecodeError:
                        continue

        # Сохранение в кэш
        if response_text:
            cache[key] = response_text
            save_cache(cache)

    except requests.exceptions.ConnectionError:
        error_msg = "⚠️ Не удалось подключиться к Ollama. Убедитесь, что сервер запущен: ollama serve"
        if stream_callback:
            stream_callback(error_msg)
        return error_msg
    except Exception as e:
        error_msg = f"⚠️ Ошибка Ollama: {str(e)}"
        if stream_callback:
            stream_callback(error_msg)
        return error_msg

    return response_text

# ================= АНАЛИЗ ЧАНКА =================
def analyze_chunk(args):
    """Анализ отдельного чанка текста"""
    i, total, chunk = args
    result = []

    def collect(text):
        result.append(text)

    prompt = f"""Ты — юрист-аналитик. Найди в фрагменте договора ТОЛЬКО риски и подозрительные моменты.
Формат ответа:
• [Риск] описание проблемы
• [Неясность] что требует уточнения
• [Выгодно] если есть плюсы для клиента

Текст для анализа:
{chunk}"""

    send_to_ollama(prompt, stream_callback=collect)
    return {"idx": i, "content": "".join(result)}

# ================= SSE ГЕНЕРАТОР =================
def generate_sse(file_path, request_id):
    """Генератор SSE-событий для стриминга результатов"""

    def send_event(data, event_type="message"):
        """Формирует корректное SSE-событие"""
        payload = {
            "type": event_type,
            "request_id": request_id,
            "timestamp": time.time()
        }
        if isinstance(data, dict):
            payload.update(data)
        else:
            payload["data"] = str(data)
        # ✅ Правильный формат SSE: data: <json>\n\n
        return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

    try:
        print(f"🚀 Начало обработки: {request_id}")

        # Проверка существования файла
        if not os.path.exists(file_path):
            print(f"❌ Файл не найден: {file_path}")
            yield send_event({"error": f"Файл не найден: {os.path.basename(file_path)}"}, "error")
            return

        yield send_event({"step": "reading", "file": os.path.basename(file_path)}, "status")
        time.sleep(0.1)

        # Извлечение текста
        ext = os.path.splitext(file_path)[1].lower()
        full_text = ""

        print(f"📋 Расширение: {ext}, размер: {os.path.getsize(file_path)} байт")

        try:
            if ext == ".pdf":
                full_text = extract_text_from_pdf(file_path)
            elif ext in [".docx", ".doc"]:
                full_text = extract_text_from_docx(file_path)
            elif ext == ".txt":
                full_text = extract_text_from_txt(file_path)
            else:
                yield send_event({
                    "error": f"Формат '{ext}' не поддерживается",
                    "supported": [".pdf", ".docx", ".doc", ".txt"]
                }, "error")
                return
        except Exception as extract_err:
            error_msg = f"Ошибка извлечения текста: {str(extract_err)}"
            print(f"❌ {error_msg}")
            yield send_event({"error": error_msg, "step": "extraction_failed"}, "error")
            return

        # Валидация извлечённого текста
        if not full_text or len(full_text.strip()) < 50:
            yield send_event({
                "error": "Текст не извлечён или слишком короткий",
                "length": len(full_text) if full_text else 0
            }, "error")
            return

        print(f"✅ Текст извлечён: {len(full_text)} символов")
        yield send_event({
            "step": "extracted",
            "chars": len(full_text),
            "words": len(full_text.split()),
            "extension": ext
        }, "status")
        time.sleep(0.1)

        # Разбивка на чанки
        chunks = split_text_smart(full_text, CHUNK_SIZE, CHUNK_OVERLAP)
        print(f"📦 Разбито на {len(chunks)} чанков")

        yield send_event({"step": "chunked", "count": len(chunks)}, "status")
        yield send_event({"step": "analyzing_start", "total": len(chunks)}, "status")

        # Параллельный анализ чанков
        tasks = [(i, len(chunks), c) for i, c in enumerate(chunks)]
        all_findings = []
        errors_count = 0

        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = {executor.submit(analyze_chunk, t): i for i, t in enumerate(tasks)}
            done = 0

            for future in as_completed(futures):
                i = futures[future]
                try:
                    res = future.result(timeout=180)  # 3 минуты на чанк
                    all_findings.append((i, res["content"]))
                    done += 1

                    preview = res["content"][:200] + ("..." if len(res["content"]) > 200 else "")

                    yield send_event({
                        "step": "chunk_done",
                        "idx": i + 1,
                        "total": len(chunks),
                        "preview": preview
                    }, "chunk_result")

                    progress = min(10 + int(done / len(chunks) * 80), 90)
                    yield send_event({
                        "progress_percent": progress,
                        "processed": done,
                        "total": len(chunks)
                    }, "progress")

                except Exception as e:
                    errors_count += 1
                    print(f"❌ Ошибка чанка {i+1}: {e}")
                    yield send_event({
                        "error": f"Ошибка части {i+1}: {str(e)[:100]}",
                        "idx": i + 1
                    }, "error")

        print(f"✅ Анализ завершён: {len(all_findings)} чанков, {errors_count} ошибок")

        # Сортировка и подготовка к синтезу
        all_findings.sort(key=lambda x: x[0])
        findings_text = [f"=== Часть {i+1} из {len(all_findings)} ===\n{t}\n"
                        for i, t in all_findings]

        yield send_event({"step": "summary_start"}, "status")

        # Финальный синтез заключения
        summary_prompt = f"""Ты — старший юрист с 20-летним опытом. Проанализируй результаты анализа договора:

{''.join(findings_text)}

Сделай СТРУКТУРИРОВАННОЕ заключение на русском языке:

🔴 ГЛАВНЫЕ РИСКИ:
1.
2.
3.
4.
5.

🟡 РЕКОМЕНДАЦИИ по изменению договора:
-

📊 МЕТРИКИ:
- Общая безопасность: __%
- Количество рисков: __
- Количество неясностей: __

⚖️ ВЕРДИКТ: [✅ Можно подписывать / ⚠️ Требует правок / ❌ Не подписывать]

Краткое обоснование вердикта (2-3 предложения):
"""

        print("🧠 Генерация финального заключения...")
        try:
            final_summary = send_to_ollama(summary_prompt)

            if final_summary and not final_summary.startswith("⚠️"):
                yield send_event({
                    "step": "summary_complete",
                    "content": final_summary,
                    "length": len(final_summary)
                }, "summary_complete")
                print("✅ Заключение сгенерировано")
            else:
                yield send_event({
                    "error": f"Не удалось сгенерировать заключение: {final_summary[:100]}",
                    "step": "summary_error"
                }, "error")
        except Exception as e:
            yield send_event({
                "error": f"Ошибка синтеза: {str(e)}",
                "step": "summary_error"
            }, "error")

        # Завершение
        yield send_event({
            "step": "complete",
            "status": "success",
            "chunks_processed": len(all_findings),
            "errors_count": errors_count
        }, "complete")

        print(f"🎉 Обработка {request_id} завершена")

    except GeneratorExit:
        print(f"⚠️ Клиент отключился: {request_id}")
        raise
    except Exception as e:
        error_detail = f"{str(e)}\n\n{traceback.format_exc()}"
        print(f"❌ Критическая ошибка SSE: {error_detail}")
        yield send_event({
            "type": "error",
            "error": error_detail[:500],  # ограничиваем длину
            "step": "critical_error"
        }, "error")
    finally:
        # Безопасное удаление файла
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"🗑️ Файл удалён: {os.path.basename(file_path)}")
        except Exception as e:
            print(f"⚠️ Не удалось удалить файл: {e}")

# ================= ЭНДПОИНТЫ =================
@app.route("/api/analyze", methods=["POST"])
def analyze_upload():
    """Загрузка файла и начало анализа"""
    try:
        print("📥 Запрос на загрузку файла")

        if "file" not in request.files:
            return jsonify({
                "error": "Нет файла в запросе",
                "hint": "Используйте FormData с полем 'file'"
            }), 400

        file = request.files["file"]

        if not file.filename or file.filename == "":
            return jsonify({"error": "Пустое имя файла"}), 400

        # Генерация уникального ID
        rid = hashlib.md5(f"{file.filename}{time.time()}{os.urandom(4).hex()}".encode()).hexdigest()[:12]
        safe_filename = "".join(c for c in file.filename if c.isalnum() or c in '._-')
        path = os.path.join(TEMP_DIR, f"{rid}_{safe_filename}")

        print(f"💾 Сохранение: {path}")
        file.save(path)

        if not os.path.exists(path):
            return jsonify({"error": "Не удалось сохранить файл"}), 500

        file_size = os.path.getsize(path)
        print(f"✅ Файл сохранён: {os.path.basename(path)} ({file_size} байт, ID: {rid})")

        # Сохранение информации о сессии
        with sessions_lock:
            active_sessions[rid] = {
                "file_path": path,
                "filename": file.filename,
                "size": file_size,
                "created_at": time.time(),
                "status": "ready"
            }

        return jsonify({
            "request_id": rid,
            "status": "ready",
            "filename": file.filename,
            "size": file_size
        })

    except Exception as e:
        error_msg = f"❌ Upload error: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        return jsonify({"error": str(e)}), 500

@app.route("/api/stream/<rid>")
def stream_sse(rid):
    """SSE-стриминг результатов анализа"""
    print(f"🔍 Запрос стрима для: {rid}")

    # Поиск файла
    fpath = None
    with sessions_lock:
        if rid in active_sessions:
            fpath = active_sessions[rid].get("file_path")
            active_sessions[rid]["status"] = "streaming"

    # Если не найдено в сессиях — ищем в папке
    if not fpath:
        try:
            for f in os.listdir(TEMP_DIR):
                if f.startswith(rid):
                    fpath = os.path.join(TEMP_DIR, f)
                    print(f"✅ Найден файл: {f}")
                    break
        except Exception as e:
            print(f"❌ Ошибка поиска: {e}")

    if not fpath or not os.path.exists(fpath):
        print(f"❌ Файл не найден для: {rid}")

        def error_stream():
            error_data = {
                "type": "error",
                "error": "Сессия не найдена или файл удалён",
                "request_id": rid,
                "timestamp": time.time()
            }
            yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"

        return Response(
            stream_with_context(error_stream()),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache, no-store, must-revalidate",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
            }
        )

    def event_stream():
        try:
            print(f"📡 Начало стриминга: {rid}")
            for event in generate_sse(fpath, rid):
                yield event
                time.sleep(0.01)  # небольшая пауза для стабильности
        except (GeneratorExit, KeyboardInterrupt):
            print(f"🔌 Клиент отключился: {rid}")
        except Exception as e:
            print(f"❌ Stream error: {e}\n{traceback.format_exc()}")
        finally:
            with sessions_lock:
                if rid in active_sessions:
                    del active_sessions[rid]

    return Response(
        stream_with_context(event_stream()),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate, no-transform",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
            "Transfer-Encoding": "chunked",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "Cache-Control",
        }
    )

@app.route("/api/status/<rid>")
def session_status(rid):
    """Проверка статуса сессии"""
    with sessions_lock:
        if rid in active_sessions:
            session = active_sessions[rid]
            return jsonify({
                "status": session.get("status", "unknown"),
                "filename": session.get("filename"),
                "size": session.get("size"),
                "created_at": session.get("created_at"),
                "file_exists": os.path.exists(session.get("file_path", ""))
            })

    # Поиск файла в папке
    for f in os.listdir(TEMP_DIR):
        if f.startswith(rid):
            fpath = os.path.join(TEMP_DIR, f)
            return jsonify({
                "status": "file_ready",
                "filename": f,
                "size": os.path.getsize(fpath),
                "file_exists": True
            })

    return jsonify({"status": "not_found", "request_id": rid}), 404

@app.route("/health")
def health():
    """Health check endpoint"""
    ollama_status = "offline"
    ollama_models = []

    try:
        r = requests.get(OLLAMA_URL.replace("/generate", "/tags"), timeout=5)
        if r.status_code == 200:
            ollama_status = "ok"
            try:
                data = r.json()
                ollama_models = [m.get("name", "") for m in data.get("models", [])]
            except:
                pass
    except Exception as e:
        ollama_status = f"error: {str(e)}"

    temp_files = 0
    temp_size = 0
    if os.path.exists(TEMP_DIR):
        try:
            files = os.listdir(TEMP_DIR)
            temp_files = len(files)
            temp_size = sum(os.path.getsize(os.path.join(TEMP_DIR, f))
                          for f in files if os.path.isfile(os.path.join(TEMP_DIR, f)))
        except:
            pass

    return jsonify({
        "status": "ok",
        "service": "LegalAI API",
        "version": "1.0.0",
        "model": MODEL_NAME,
        "ollama": {
            "status": ollama_status,
            "models": ollama_models,
            "url": OLLAMA_URL
        },
        "temp_dir": {
            "path": os.path.abspath(TEMP_DIR),
            "files": temp_files,
            "size_bytes": temp_size
        },
        "active_sessions": len(active_sessions),
        "timestamp": time.time()
    })

@app.route("/api/cleanup", methods=["POST"])
def cleanup():
    """Очистка старых файлов"""
    try:
        max_age = request.json.get("max_age_seconds", 3600) if request.is_json else 3600
        now = time.time()
        deleted = 0

        with sessions_lock:
            to_delete = [
                rid for rid, session in active_sessions.items()
                if now - session.get("created_at", 0) > max_age
            ]
            for rid in to_delete:
                fpath = active_sessions[rid].get("file_path")
                if fpath and os.path.exists(fpath):
                    os.remove(fpath)
                    print(f"🗑️ Удалён старый файл: {os.path.basename(fpath)}")
                del active_sessions[rid]
                deleted += 1

        return jsonify({
            "status": "ok",
            "deleted_sessions": deleted,
            "remaining_sessions": len(active_sessions)
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ================= ЗАПУСК =================
if __name__ == "__main__":
    print("=" * 60)
    print("🚀 LegalAI Python API v1.0")
    print("=" * 60)
    print(f"📦 Модель: {MODEL_NAME}")
    print(f"🌐 Ollama: {OLLAMA_URL}")
    print(f"📁 Temp: {os.path.abspath(TEMP_DIR)}")
    print(f"👷 Workers: {MAX_WORKERS}")
    print(f"📊 Chunk: {CHUNK_SIZE} (overlap: {CHUNK_OVERLAP})")
    print("=" * 60)
    print("Запуск на http://0.0.0.0:5000")
    print("Ctrl+C для остановки")
    print("=" * 60)

    # Проверка Ollama
    try:
        print("🔍 Проверка Ollama...")
        r = requests.get(OLLAMA_URL.replace("/generate", "/tags"), timeout=5)
        if r.status_code == 200:
            print("✅ Ollama подключен")
        else:
            print(f"⚠️ Ollama: код {r.status_code}")
    except Exception as e:
        print(f"❌ Ollama недоступен: {e}")
        print("💡 Запустите: ollama serve")

    print()
    app.run(host="0.0.0.0", port=5000, debug=False, threaded=True)
